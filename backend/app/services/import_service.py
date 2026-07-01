"""TripImportService — parse Excel/JSON trip files and import into DB."""

import io
import json
from typing import Optional

from sqlalchemy.ext.asyncio import AsyncSession

from app.models.trip import Trip, TripSchedule


class TripImportService:
    """Parse and import trip files."""

    @staticmethod
    def generate_template_bytes() -> bytes:
        """Generate an Excel import template."""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            wb = openpyxl.Workbook()

            # ── Sheet 1: 行程概要 ──
            ws1 = wb.active
            ws1.title = "行程概要"
            header_font = Font(bold=True, size=11, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            thin_border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin'),
            )

            fields = [
                ("行程编号", "YN-2026-001", "唯一标识，如 YN-2026-001"),
                ("行程标题", "云南丽江古城5日深度游", "行程名称"),
                ("副标题", "遇见古城，邂逅纳西文化", "可选"),
                ("目的地", "云南丽江", "主要目的地"),
                ("详细目的地", "丽江古城,玉龙雪山,泸沽湖", "逗号分隔"),
                ("国家", "中国", "中国/日本/等"),
                ("省份", "云南", ""),
                ("城市", "丽江", ""),
                ("类型", "国内游", "国内游/出境游/周边游/定制游"),
                ("天数", "5", "整数"),
                ("晚数", "4", ""),
                ("出发城市", "北京", ""),
                ("最佳季节", "春季/秋季", ""),
                ("成人价", "3980", "元/人"),
                ("儿童价", "2680", "元/人"),
                ("婴儿价", "800", "元/人"),
                ("单房差", "800", "元"),
                ("费用包含", "往返机票,酒店,门票,导游", "逗号分隔"),
                ("费用不含", "个人消费,自费项目", "逗号分隔"),
                ("成团最小人数", "1", ""),
                ("成团最大人数", "30", ""),
                ("发团日期", "2026-07-15,2026-08-01", "逗号分隔的日期"),
            ]

            for i, (field, example, note) in enumerate(fields, 1):
                ws1.cell(row=i, column=1, value=field).font = Font(bold=True)
                ws1.cell(row=i, column=2, value=example)
                ws1.cell(row=i, column=3, value=note).font = Font(color="808080", italic=True)
                for col in range(1, 4):
                    ws1.cell(row=i, column=col).border = thin_border

            ws1.column_dimensions['A'].width = 18
            ws1.column_dimensions['B'].width = 35
            ws1.column_dimensions['C'].width = 30

            # ── Sheet 2: 行程安排 ──
            ws2 = wb.create_sheet("行程安排")
            schedule_headers = ["天数", "主题", "类型", "地点", "活动内容", "描述", "餐饮", "酒店名称", "酒店星级", "交通方式", "交通详情", "注意事项"]
            example_row = ["1", "出发抵达，初识丽江", "交通", "丽江三义机场", "接机入住酒店", "专车接机，入住四星级酒店，晚间自由逛古城", "晚餐", "丽江古城酒店", "4", "飞机+专车", "航班约3小时", "注意高原反应，多喝水"]

            for i, h in enumerate(schedule_headers, 1):
                cell = ws2.cell(row=1, column=i, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
                ws2.column_dimensions[get_column_letter(i)].width = 14

            for i, val in enumerate(example_row, 1):
                ws2.cell(row=2, column=i, value=val).border = thin_border

            buf = io.BytesIO()
            wb.save(buf)
            return buf.getvalue()
        except ImportError:
            raise RuntimeError("openpyxl 未安装，无法生成模板")

    @staticmethod
    async def parse_and_import(
        content: bytes,
        filename: str,
        db: AsyncSession,
        created_by: str,
    ) -> dict:
        """Parse uploaded file and import into DB.

        Returns dict with: trip_id, trip_title, schedules_count, warnings
        """
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        warnings = []

        if ext == "xlsx":
            parsed = TripImportService._parse_excel(content)
        elif ext == "json":
            parsed = json.loads(content.decode("utf-8"))
            parsed = {
                "trip": parsed.get("trip", parsed),
                "schedules": parsed.get("schedules", parsed.get("itinerary", []))}
        elif ext == "docx":
            parsed = TripImportService._parse_docx(content)
        elif ext == "pdf":
            parsed = TripImportService._parse_pdf(content)
        else:
            raise ValueError(f"不支持的文件格式: .{ext}，请上传 .xlsx / .json / .docx / .pdf")

        trip_data = parsed.get("trip", {})
        schedules_data = parsed.get("schedules", [])

        if not trip_data.get("title") and not trip_data.get("destination"):
            raise ValueError("无法解析行程信息：缺少标题或目的地")

        # Check duplicate code
        code = trip_data.get("code", "")
        if code:
            from sqlalchemy import select as sel
            existing = await db.execute(sel(Trip).where(Trip.code == code))
            if existing.scalar_one_or_none():
                code = f"{code}_imported"
                warnings.append(f"行程编号已存在，自动改为 {code}")

        if not code:
            code = await TripImportService._generate_code(trip_data, db)

        # Parse numeric fields
        def _int(v, default=0):
            try:
                return int(v) if v else default
            except (ValueError, TypeError):
                return default

        def _float(v, default=None):
            try:
                return float(v) if v else default
            except (ValueError, TypeError):
                return default

        def _list(v, default=None):
            if default is None:
                default = []
            if isinstance(v, list):
                return v
            if isinstance(v, str) and v.strip():
                return [x.strip() for x in v.split(",")]
            return default

        # Create trip
        trip = Trip(
            code=code,
            title=str(trip_data.get("title", "未命名行程")),
            subtitle=str(trip_data.get("subtitle", "")) if trip_data.get("subtitle") else None,
            destination=str(trip_data.get("destination", "")),
            destinations_detail=_list(trip_data.get("destinations_detail", trip_data.get("详细目的地"))),
            country=str(trip_data.get("country", "中国")),
            province=str(trip_data.get("province", "")) if trip_data.get("province") else None,
            city=str(trip_data.get("city", "")) if trip_data.get("city") else None,
            category=str(trip_data.get("category", trip_data.get("类型", "国内游"))),
            duration_days=_int(trip_data.get("duration_days", trip_data.get("天数", 1))),
            duration_nights=_int(trip_data.get("duration_nights", trip_data.get("晚数"))),
            departure_city=str(trip_data.get("departure_city", trip_data.get("出发城市", ""))) if trip_data.get("departure_city") else None,
            best_season=str(trip_data.get("best_season", trip_data.get("最佳季节", ""))) if trip_data.get("best_season") else None,
            price_adult=_float(trip_data.get("price_adult", trip_data.get("成人价"))),
            price_child=_float(trip_data.get("price_child", trip_data.get("儿童价"))),
            price_infant=_float(trip_data.get("price_infant", trip_data.get("婴儿价"))),
            single_room_supplement=_float(trip_data.get("single_room_supplement", trip_data.get("单房差"))),
            price_includes=_list(trip_data.get("price_includes", trip_data.get("费用包含"))),
            price_excludes=_list(trip_data.get("price_excludes", trip_data.get("费用不含"))),
            group_size_min=_int(trip_data.get("group_size_min", trip_data.get("成团最小人数", 1))),
            group_size_max=_int(trip_data.get("group_size_max", trip_data.get("成团最大人数"))) or None,
            departure_dates=_list(trip_data.get("departure_dates", trip_data.get("发团日期"))),
            summary=str(trip_data.get("summary", "")) if trip_data.get("summary") else None,
            highlights=_list(trip_data.get("highlights", [])),
            recommendation_reasons=_list(trip_data.get("recommendation_reasons", [])),
            detailed_description=str(trip_data.get("detailed_description", "")) if trip_data.get("detailed_description") else None,
            status="active",
            created_by=created_by,
        )
        db.add(trip)
        await db.flush()

        # Create schedules
        schedule_count = 0
        for s in schedules_data:
            day = _int(s.get("day_number", s.get("天数", s.get("day", 0))))
            if day <= 0:
                continue
            schedule = TripSchedule(
                trip_id=trip.id,
                day_number=day,
                theme=str(s.get("theme", s.get("主题", ""))) if s.get("theme") or s.get("主题") else None,
                schedule_type=str(s.get("schedule_type", s.get("类型", "景点"))),
                location=str(s.get("location", s.get("地点", ""))) if s.get("location") or s.get("地点") else None,
                activity=str(s.get("activity", s.get("活动内容", ""))) if s.get("activity") or s.get("活动内容") else None,
                description=str(s.get("description", s.get("描述", ""))) if s.get("description") or s.get("描述") else None,
                meal_included=str(s.get("meal_included", s.get("餐饮", ""))) if s.get("meal_included") or s.get("餐饮") else None,
                hotel_name=str(s.get("hotel_name", s.get("酒店名称", ""))) if s.get("hotel_name") or s.get("酒店名称") else None,
                hotel_stars=_int(s.get("hotel_stars", s.get("酒店星级"))) or None,
                transport_type=str(s.get("transport_type", s.get("交通方式", ""))) if s.get("transport_type") or s.get("交通方式") else None,
                transport_detail=str(s.get("transport_detail", s.get("交通详情", ""))) if s.get("transport_detail") or s.get("交通详情") else None,
                tips=str(s.get("tips", s.get("注意事项", ""))) if s.get("tips") or s.get("注意事项") else None,
                sort_order=schedule_count,
            )
            db.add(schedule)
            schedule_count += 1

        await db.flush()

        return {
            "trip_id": str(trip.id),
            "trip_title": trip.title,
            "schedules_count": schedule_count,
            "validation_warnings": warnings,
        }

    # ── Province/City abbreviation mapping for code generation ──
    PROVINCE_ABBR = {
        "新疆": "XJ", "西藏": "XZ", "云南": "YN", "海南": "HN", "四川": "SC",
        "贵州": "GZ", "广西": "GX", "广东": "GD", "福建": "FJ", "浙江": "ZJ",
        "江苏": "JS", "上海": "SH", "北京": "BJ", "天津": "TJ", "重庆": "CQ",
        "山东": "SD", "河北": "HB", "河南": "HA", "湖北": "HB2", "湖南": "HN2",
        "江西": "JX", "安徽": "AH", "陕西": "SN", "甘肃": "GS", "青海": "QH",
        "宁夏": "NX", "内蒙古": "NM", "辽宁": "LN", "吉林": "JL", "黑龙江": "HL",
        "日本": "JP", "泰国": "TH", "韩国": "KR", "越南": "VN", "欧洲": "EU",
    }
    CITY_ABBR = {
        "伊犁": "YL", "乌鲁木齐": "WQ", "丽江": "LJ", "三亚": "SY", "东京": "DJ",
        "京都": "JD", "大阪": "DB", "独库": "DK", "天山": "TS", "赛里木湖": "SL",
        "那拉提": "NL", "吐鲁番": "TL", "喀纳斯": "KN", "大理": "DL", "昆明": "KM",
        "海口": "HK", "成都": "CD", "重庆": "CQ", "西安": "XA", "青岛": "QD",
        "敦煌": "DH", "桂林": "GL", "张家界": "ZJ", "黄山": "HS", "厦门": "XM",
        "呼伦贝尔": "HL", "长白山": "CB", "九寨沟": "JZ", "稻城": "DC",
    }

    @staticmethod
    async def _generate_code(trip_data: dict, db) -> str:
        """Generate standard trip code: {province}-{city}-{year}-{seq}.

        Matches seed data format: YN-LJ-2026-001, HN-SY-2026-002, etc.
        """
        from datetime import datetime
        from sqlalchemy import select, func
        from app.models.trip import Trip

        dest = trip_data.get("destination", "")
        province = trip_data.get("province", "")
        city = trip_data.get("city", "")
        title = trip_data.get("title", "")
        all_text = f"{title} {dest} {province} {city}"

        # Find province abbreviation
        p_abbr = ""
        for name, abbr in TripImportService.PROVINCE_ABBR.items():
            if name in all_text:
                p_abbr = abbr
                break
        if not p_abbr:
            p_abbr = "CN"

        # Find city abbreviation
        c_abbr = ""
        for name, abbr in TripImportService.CITY_ABBR.items():
            if name in all_text:
                c_abbr = abbr
                break
        if not c_abbr:
            # Use first 2 chars of destination as fallback
            import re
            c_abbr = re.sub(r'[^A-Za-z]', '', dest[:4].upper())[:2] or "XX"

        year = str(datetime.now().year)

        # Find next sequence number
        prefix = f"{p_abbr}-{c_abbr}-{year}-"
        result = await db.execute(
            select(Trip.code).where(Trip.code.like(f"{prefix}%"))
        )
        existing = result.scalars().all()
        seq = 1
        if existing:
            nums = []
            for code in existing:
                try:
                    nums.append(int(code.split("-")[-1]))
                except (ValueError, IndexError):
                    pass
            if nums:
                seq = max(nums) + 1

        return f"{prefix}{seq:03d}"

    @staticmethod
    def _parse_excel(content: bytes) -> dict:
        """Parse Excel file into trip + schedules dict."""
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)

        trip_data = {}
        schedules = []

        # Parse overview sheet
        for sheet_name in ["行程概要", "概览", "Sheet1"]:
            if sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows(min_row=1, values_only=True):
                    if row[0] and row[1]:
                        key = str(row[0]).strip()
                        val = row[1]
                        if key and val:
                            field_map = {
                                "行程编号": "code", "行程标题": "title", "副标题": "subtitle",
                                "目的地": "destination", "详细目的地": "destinations_detail",
                                "国家": "country", "省份": "province", "城市": "city",
                                "类型": "category", "天数": "duration_days", "晚数": "duration_nights",
                                "出发城市": "departure_city", "最佳季节": "best_season",
                                "成人价": "price_adult", "儿童价": "price_child", "婴儿价": "price_infant",
                                "单房差": "single_room_supplement",
                                "费用包含": "price_includes", "费用不含": "price_excludes",
                                "成团最小人数": "group_size_min", "成团最大人数": "group_size_max",
                                "发团日期": "departure_dates",
                            }
                            field = field_map.get(key, key)
                            trip_data[field] = val
                break

        # Parse schedule sheet
        for sheet_name in ["行程安排", "日程", "行程表"]:
            if sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows(min_row=1, values_only=True))
                if not rows:
                    continue
                headers = [str(h).strip() if h else "" for h in rows[0]]
                header_map = {
                    "天数": "day_number", "天": "day_number", "day": "day_number",
                    "主题": "theme", "类型": "schedule_type", "type": "schedule_type",
                    "地点": "location", "活动内容": "activity", "活动": "activity",
                    "描述": "description", "内容": "description",
                    "餐饮": "meal_included", "餐": "meal_included",
                    "酒店名称": "hotel_name", "酒店": "hotel_name", "住宿": "hotel_name",
                    "酒店星级": "hotel_stars", "星级": "hotel_stars",
                    "交通方式": "transport_type", "交通": "transport_type",
                    "交通详情": "transport_detail",
                    "注意事项": "tips", "注意": "tips",
                }
                for row in rows[1:]:
                    if not any(row):
                        continue
                    schedule = {}
                    for i, cell in enumerate(row):
                        if i >= len(headers) or cell is None:
                            continue
                        h = headers[i].lower() if headers[i] else ""
                        field = header_map.get(headers[i], header_map.get(h, h))
                        schedule[field] = str(cell).strip() if cell else ""
                    if schedule.get("day_number"):
                        try:
                            schedule["day_number"] = int(str(schedule["day_number"]).replace("第", "").replace("天", ""))
                        except ValueError:
                            schedule["day_number"] = len(schedules) + 1
                    else:
                        schedule["day_number"] = len(schedules) + 1
                    schedules.append(schedule)
                break

        return {"trip": trip_data, "schedules": schedules}

    @staticmethod
    def _parse_docx(content: bytes) -> dict:
        """Parse DOCX travel itinerary file into structured trip + schedules."""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx 未安装，无法解析 DOCX 文件")

        doc = Document(io.BytesIO(content))

        # ── Extract all paragraphs ──
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                # Detect bold/heading style
                is_bold = any(run.bold for run in p.runs if run.bold)
                style = p.style.name if p.style else ""
                paragraphs.append({"text": text, "bold": is_bold, "style": style})

        # ── Extract all tables ──
        tables = []
        for t in doc.tables:
            rows = []
            for row in t.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows.append(cells)
            tables.append(rows)

        trip = {}
        schedules = []
        import re

        # ── Parse title from first meaningful paragraphs ──
        title_parts = []
        for p in paragraphs[:6]:
            txt = p["text"]
            # Skip short promotional lines like "真纯玩不擦边"
            if len(txt) > 6 and ("伊犁" in txt or "新疆" in txt or "双飞" in txt or "日" in txt):
                title_parts.append(txt)
        trip["title"] = title_parts[0] if title_parts else paragraphs[0]["text"] if paragraphs else "未命名行程"
        trip["subtitle"] = title_parts[1] if len(title_parts) > 1 else ""

        # ── Extract price ──
        for p in paragraphs:
            m = re.search(r'(?:成人|价格)[：:]?\s*(\d{3,5})\s*元', p["text"])
            if m:
                trip["price_adult"] = float(m.group(1))
                break
        # Also check table cells
        for t in tables:
            for row in t:
                row_text = " ".join(row)
                m = re.search(r'成人[：:]?\s*(\d{3,5})\s*元', row_text)
                if m:
                    trip["price_adult"] = float(m.group(1))
                    break

        # ── Extract children's price ──
        for p in paragraphs:
            m = re.search(r'(?:小孩|儿童|12周岁以下)[^\d]*(\d{3,5})\s*元', p["text"])
            if m:
                trip["price_child"] = float(m.group(1))
                break
        if "price_child" not in trip:
            for t in tables:
                for row in t:
                    row_text = " ".join(row)
                    m = re.search(r'(?:小孩|儿童|12周岁以下)[^\d]*(\d{3,5})\s*元', row_text)
                    if m:
                        trip["price_child"] = float(m.group(1))
                        break

        # ── Extract duration ──
        for p in paragraphs:
            m = re.search(r'(?:双飞|纯玩)?(\d+)\s*日', p["text"])
            if m and int(m.group(1)) >= 2:
                trip["duration_days"] = int(m.group(1))
                break
        if "duration_days" not in trip:
            trip["duration_days"] = 8  # default

        # ── Extract destination ──
        dest_keywords = ["新疆", "伊犁", "独库", "天山", "乌鲁木齐", "赛里木湖", "那拉提", "吐鲁番", "喀纳斯"]
        found_dests = []
        all_text = " ".join(p["text"] for p in paragraphs)
        for kw in dest_keywords:
            if kw in all_text:
                found_dests.append(kw)
        trip["destination"] = found_dests[0] if found_dests else "新疆"
        trip["destinations_detail"] = found_dests[:5] if found_dests else []

        # ── Extract departure city ──
        for p in paragraphs:
            m = re.search(r'(?:集合出发|出发).*?(?:桐乡|嘉兴|杭州|上海|北京|广州|深圳|成都|南京|武汉)', p["text"])
            if m:
                # Extract just the city name
                city_m = re.search(r'(桐乡|嘉兴|杭州|上海|北京|广州|深圳|成都|南京|武汉)', p["text"])
                if city_m:
                    trip["departure_city"] = city_m.group(1)
                    break
        # Also check overall text
        if "departure_city" not in trip:
            all_paras = " ".join(p["text"] for p in paragraphs)
            m = re.search(r'(?:桐乡|嘉兴|杭州).*?(?:集合出发|出发)', all_paras)
            if m:
                m2 = re.search(r'(桐乡|嘉兴|杭州)', all_paras)
                if m2:
                    trip["departure_city"] = m2.group(1)
        trip["category"] = "国内游"
        trip["country"] = "中国"
        trip["province"] = "新疆"

        # ── Parse itinerary table (quick reference table) ──
        # Find table with day columns (DAY1, D1, 第1天 等)
        overview_table = None
        for t in tables:
            if len(t) >= 3 and t[0]:
                first_row = " ".join(t[0]).lower()
                if any(kw in first_row for kw in ["day", "天", "行程", "用餐", "住宿"]):
                    overview_table = t
                    break
        if not overview_table and tables:
            overview_table = tables[0]

        if overview_table:
            # Determine which columns contain what
            header = overview_table[0]
            day_col = None
            route_col = None
            meal_col = None
            hotel_col = None
            for i, h in enumerate(header):
                h_l = h.lower()
                if any(kw in h_l for kw in ["day", "天", "d"]):
                    day_col = i
                elif any(kw in h_l for kw in ["行程", "路线", "安排"]):
                    route_col = i
                elif any(kw in h_l for kw in ["餐", "吃", "用餐"]):
                    meal_col = i
                elif any(kw in h_l for kw in ["住宿", "酒店", "住"]):
                    hotel_col = i

            for row in overview_table[1:]:
                if not any(row):
                    continue
                day_text = str(row[day_col]).strip() if day_col is not None and day_col < len(row) else ""
                m = re.search(r'(\d+)', day_text)
                day_num = int(m.group(1)) if m else len(schedules) + 1
                if day_num < 1 or day_num > 15:
                    continue

                route = str(row[route_col]).strip() if route_col is not None and route_col < len(row) else ""
                meal = str(row[meal_col]).strip() if meal_col is not None and meal_col < len(row) else ""
                hotel = str(row[hotel_col]).strip() if hotel_col is not None and hotel_col < len(row) else ""

                schedules.append({
                    "day_number": day_num,
                    "theme": route,
                    "schedule_type": "景点",
                    "meal_included": meal,
                    "hotel_name": hotel,
                    "description": "",
                    "tips": "",
                })

        # ── Parse detailed day descriptions from paragraphs ──
        day_pattern = re.compile(r'(?:DAY|D|第)\s*(\d{1,2})\s*(?:天|日)', re.IGNORECASE)
        current_day = None
        day_buffer = []
        tips_buffer = []

        for p in paragraphs:
            txt = p["text"]
            dm = day_pattern.search(txt)
            if dm:
                # Save previous day's content
                if current_day is not None and day_buffer:
                    for s in schedules:
                        if s["day_number"] == current_day:
                            s["description"] = " ".join(day_buffer)
                            if tips_buffer:
                                s["tips"] = " ".join(tips_buffer)
                            break
                    else:
                        # Day not in overview table: add it
                        schedules.append({
                            "day_number": current_day,
                            "theme": "",
                            "schedule_type": "景点",
                            "description": " ".join(day_buffer),
                            "tips": " ".join(tips_buffer) if tips_buffer else "",
                        })
                current_day = int(dm.group(1))
                day_buffer = [txt]
                tips_buffer = []
            elif current_day is not None:
                if "温馨提示" in txt or "注意" in txt:
                    tips_buffer.append(txt)
                elif "住宿" in txt or "酒店" in txt:
                    # Extract hotel info
                    for s in schedules:
                        if s["day_number"] == current_day and not s.get("hotel_name"):
                            s["hotel_name"] = txt
                    day_buffer.append(txt)
                else:
                    day_buffer.append(txt)

        # Save last day
        if current_day is not None and day_buffer:
            for s in schedules:
                if s["day_number"] == current_day:
                    s["description"] = " ".join(day_buffer)
                    if tips_buffer:
                        s["tips"] = " ".join(tips_buffer)
                    break
            else:
                schedules.append({
                    "day_number": current_day,
                    "theme": "",
                    "schedule_type": "景点",
                    "description": " ".join(day_buffer),
                    "tips": " ".join(tips_buffer) if tips_buffer else "",
                })

        # ── Parse fee sections ──
        includes = []
        excludes = []
        in_excludes = False
        for p in paragraphs:
            txt = p["text"]
            if any(kw in txt for kw in ["费用不含", "不含", "费用不包含"]):
                in_excludes = True
                continue
            if any(kw in txt for kw in ["费用包含", "包含项目"]):
                in_excludes = False
                continue

            # Match list items: "1. xxx" or "- xxx" or "xxx：yyy"
            items = re.findall(r'(?:^\d+[\.\、)]\s*|[-•]\s*)(.+)', txt)
            if items:
                if in_excludes:
                    excludes.extend(items)
                else:
                    includes.extend(items)

        # Fallback: parse tables for fees
        if not includes:
            for t in tables:
                for row in t:
                    row_text = " ".join(row)
                    if "费用包含" in row_text or "包含" in row_text:
                        includes.extend([c for c in row if c and "费用" not in c and len(c) > 5])
                    if "不含" in row_text:
                        excludes.extend([c for c in row if c and "费用" not in c and len(c) > 5])

        trip["price_includes"] = includes[:10] if includes else []
        trip["price_excludes"] = excludes[:10] if excludes else []

        # ── Extract highlights ──
        highlights = []
        # Patterns: bullet points after "亮点"/"特色"/"卖点" keywords
        in_highlights = False
        for p in paragraphs:
            txt = p["text"]
            if any(kw in txt for kw in ["产品亮点", "行程亮点", "线路特色", "推荐理由", "产品特色"]):
                in_highlights = True
                continue
            if in_highlights and len(txt) > 8:
                if any(kw in txt for kw in ["费用包含", "费用不含", "温馨提示", "注意事项", "报名须知"]):
                    in_highlights = False
                    continue
                # Clean up bullet markers
                cleaned = re.sub(r'^[\d\.\、\-\•\★\☆\s]+', '', txt)
                if len(cleaned) > 5:
                    highlights.append(cleaned)
                if len(highlights) >= 6:
                    break
        # Fallback: look for "0购物0自费" style selling points
        if not highlights:
            for p in paragraphs:
                txt = p["text"]
                if any(kw in txt for kw in ["0购物", "纯玩", "一价全含", "自组团", "管家", "赠送", "升级"]):
                    highlights.append(txt[:80])
                    if len(highlights) >= 5:
                        break
        trip["highlights"] = highlights[:6]

        # ── Auto-generate summary ──
        dest = trip.get("destination", "")
        days = trip.get("duration_days", "")
        price = trip.get("price_adult", "")
        departure = trip.get("departure_city", "")
        title = trip.get("title", "")

        summary_parts = [f"{dest}{days}日游"]
        if "纯玩" in title:
            summary_parts.append("纯玩0购物")
        if price:
            summary_parts.append(f"成人¥{int(price)}/人起")
        if departure:
            summary_parts.append(f"从{departure}出发")
        if highlights:
            summary_parts.append(f"特色：{'、'.join(highlights[:3])}")

        trip["summary"] = f"{title}，{'，'.join(summary_parts[1:])}。" if len(summary_parts) > 1 else title

        # ── Generate recommendation reasons ──
        reasons = []
        if "纯玩" in title or "0购物" in title:
            reasons.append("纯玩0购物，品质保障")
        if price:
            reasons.append(f"超高性价比，成人仅需¥{int(price)}/人")
        if departure:
            reasons.append(f"{departure}出发，交通便利")
        if highlights:
            reasons.append(highlights[0] if highlights[0] else "")
        trip["recommendation_reasons"] = [r for r in reasons if r][:3]

        return {"trip": trip, "schedules": schedules}

    @staticmethod
    def _parse_pdf(content: bytes) -> dict:
        """Parse PDF travel itinerary into structured trip + schedules."""
        all_text = ""
        tables = []

        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        all_text += text + "\n"
                    page_tables = page.extract_tables()
                    for t in page_tables:
                        rows = []
                        for row in t:
                            cells = [str(c).strip() if c else "" for c in row]
                            rows.append(cells)
                        if rows:
                            tables.append(rows)
        except Exception:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(content))
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        all_text += text + "\n"
            except Exception:
                all_text = "[PDF parsing requires pdfplumber or PyPDF2]"

        import re
        trip = {"country": "中国", "category": "国内游"}
        schedules = []
        lines = all_text.split("\n")

        # ── Title ──
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 6:
                trip["title"] = line
                break
        if "title" not in trip:
            trip["title"] = "未命名行程"

        # ── Price ──
        for line in lines:
            m = re.search(r'(?:成人|价格)[：:]?\s*(\d{3,5})\s*元', line)
            if m:
                trip["price_adult"] = float(m.group(1))
                break

        # ── Duration ──
        for line in lines:
            m = re.search(r'(\d+)\s*日', line)
            if m and int(m.group(1)) >= 2:
                trip["duration_days"] = int(m.group(1))
                break
        if "duration_days" not in trip:
            trip["duration_days"] = 8

        # ── Destination ──
        dest_keywords = ["新疆", "伊犁", "独库", "天山", "赛里木湖", "那拉提", "吐鲁番", "喀纳斯"]
        found = [kw for kw in dest_keywords if kw in all_text]
        trip["destination"] = found[0] if found else "未知"
        trip["destinations_detail"] = found[:5] if found else []

        # ── Parse tables for schedules ──
        for t in tables:
            if len(t) >= 3:
                header = " ".join(t[0]).lower()
                if any(kw in header for kw in ["day", "天", "行程", "用餐"]):
                    day_col = route_col = meal_col = hotel_col = None
                    for i, h in enumerate(t[0]):
                        h_l = h.lower()
                        if any(kw in h_l for kw in ["day", "天"]):
                            day_col = i
                        elif any(kw in h_l for kw in ["行程", "安排"]):
                            route_col = i
                        elif any(kw in h_l for kw in ["餐"]):
                            meal_col = i
                        elif any(kw in h_l for kw in ["住宿", "酒店"]):
                            hotel_col = i
                    for row in t[1:]:
                        day_text = row[day_col] if day_col is not None and day_col < len(row) else ""
                        m = re.search(r'(\d+)', day_text)
                        if m:
                            schedules.append({
                                "day_number": int(m.group(1)),
                                "theme": row[route_col] if route_col is not None and route_col < len(row) else "",
                                "schedule_type": "景点",
                                "meal_included": row[meal_col] if meal_col is not None and meal_col < len(row) else "",
                                "hotel_name": row[hotel_col] if hotel_col is not None and hotel_col < len(row) else "",
                                "description": "",
                            })

        # ── Parse day descriptions from text ──
        if not schedules:
            day_pattern = re.compile(r'(?:第|D|DAY)\s*(\d{1,2})\s*(?:天|日)', re.IGNORECASE)
            current_day = None
            buffer = []
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                dm = day_pattern.search(line)
                if dm:
                    if current_day and buffer:
                        schedules.append({
                            "day_number": current_day,
                            "theme": buffer[0] if buffer else "",
                            "schedule_type": "景点",
                            "description": " ".join(buffer),
                        })
                    current_day = int(dm.group(1))
                    buffer = [line]
                elif current_day:
                    buffer.append(line)
            if current_day and buffer:
                schedules.append({
                    "day_number": current_day,
                    "theme": buffer[0] if buffer else "",
                    "schedule_type": "景点",
                    "description": " ".join(buffer),
                })

        return {"trip": trip, "schedules": schedules}
